from docx import Document

def create_document(dane):

    
    dane_kredyt = dane['tech_data']['form_data']

    document = Document()


    document.add_heading('Dane o kredycie', level=3)

    document.add_paragraph(
        f"Kwota kredytu: {round(float(dane_kredyt['kapital1']),2):,.2f} zł", style='List Bullet'
    )
    document.add_paragraph(
        f"Data uruchomienia kredytu: {dane_kredyt['dataStart1']} " , style='List Bullet'
    )
    document.add_paragraph(
        f"Liczba miesięcy kredytowania: {dane_kredyt['okresy']}" , style='List Bullet'
    )

    document.add_paragraph(
        f"Marża banku: {dane_kredyt['marza']}%" , style='List Bullet'
    )

    document.add_paragraph(
        f"Rodzaj wiboru: Wibor {dane_kredyt['rodzajWiboru']}" , style='List Bullet'
    )

    document.add_paragraph(
        f"Wibor z dnia uruchomienia kredytu: {dane['fin_data']['wibor_start']}%" , style='List Bullet'
    )

    document.add_heading('Zamrożenie wiboru', level=3)

    document.add_paragraph(
        f"Data zamrożenia wiboru: {dane_kredyt['dataZamrozenia']}" , style='List Bullet'
    )

    document.add_paragraph(
        f"Do zwrotu przez bank (raty przed zamrożeniem): {dane['do_zwrotu_przed']}" , style='List Bullet'
    )

    
    document.add_paragraph(
        f"Saldo kredytu: {dane['kapital_do_splaty_po_zamr']}" , style='List Bullet'
    )

    return document