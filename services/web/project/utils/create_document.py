from docx import Document
from docx.shared import Inches
import base64
from io import BytesIO, StringIO
from svglib.svglib import svg2rlg
from reportlab.graphics import renderPM
import xml.etree.ElementTree as ET
def create_document(dane):

    
    dane_kredyt = dane['tech_data']['form_data']

    dane_svg = dane['svg_dane1']
    print(dane_svg)

    document = Document()
    # define the SVG string
    svg_string = '''<svg class="svg-holder" width="1110" height="450" style="border: 1px solid black; background: red;"><g transform="translate(80,60)"><g class="xAxis" transform="translate(0,320)" fill="none" font-size="10" font-family="sans-serif" text-anchor="middle"><path class="domain" stroke="#000" d="M0.5,6V0.5H1000.5V6"></path><g class="tick" opacity="1" transform="translate(0.5,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="text-anchor: end; font-size: 13px;">0 zł</text></g><g class="tick" opacity="1" transform="translate(78.86014164379195,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="text-anchor: end; font-size: 13px;">50000 zł</text></g><g class="tick" opacity="1" transform="translate(157.2202832875839,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="text-anchor: end; font-size: 13px;">100000 zł</text></g><g class="tick" opacity="1" transform="translate(235.58042493137586,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="text-anchor: end; font-size: 13px;">150000 zł</text></g><g class="tick" opacity="1" transform="translate(313.9405665751678,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="text-anchor: end; font-size: 13px;">200000 zł</text></g><g class="tick" opacity="1" transform="translate(392.30070821895976,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="text-anchor: end; font-size: 13px;">250000 zł</text></g><g class="tick" opacity="1" transform="translate(470.6608498627517,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="text-anchor: end; font-size: 13px;">300000 zł</text></g><g class="tick" opacity="1" transform="translate(549.0209915065436,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="text-anchor: end; font-size: 13px;">350000 zł</text></g><g class="tick" opacity="1" transform="translate(627.3811331503356,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="text-anchor: end; font-size: 13px;">400000 zł</text></g><g class="tick" opacity="1" transform="translate(705.7412747941275,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="text-anchor: end; font-size: 13px;">450000 zł</text></g><g class="tick" opacity="1" transform="translate(784.1014164379195,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="text-anchor: end; font-size: 13px;">500000 zł</text></g><g class="tick" opacity="1" transform="translate(862.4615580817115,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="text-anchor: end; font-size: 13px;">550000 zł</text></g><g class="tick" opacity="1" transform="translate(940.8216997255034,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="text-anchor: end; font-size: 13px;">600000 zł</text></g></g><rect class="bar" x="0" y="186.66666666666666" height="79.99999999999999" width="1000"></rect><rect class="bar" x="0" y="53.33333333333334" height="79.99999999999999" width="26.36354874275068"></rect><g class="yAxis" fill="none" font-size="10" font-family="sans-serif" text-anchor="end" style="font-size: 13px;"><path class="domain" stroke="#000" d="M-6,320.5H0.5V0.5H-6"></path><g class="tick" opacity="1" transform="translate(0,226.66666666666666)"><line stroke="#000" x2="-6"></line><text fill="#000" x="-9" dy="0.32em" transform="translate(25,0)" font-size="19" rotate="0" style="text-anchor: start;">odseżąśki z tytułu wiboru zmiennego</text></g><g class="tick" opacity="1" transform="translate(0,93.33333333333334)"><line stroke="#000" x2="-6"></line><text fill="#000" x="-9" dy="0.32em" transform="translate(25,0)" font-size="19" rotate="0" style="text-anchor: start;">odsetki z tytułu wiboru stałego</text></g></g></g></svg>'''
    svg_string2 = """<svg class="svg-holder" width="1110" height="450" style="border: 1px solid black; background: white;"><g transform="translate(80,60)"><g class="xAxis" transform="translate(0,320)" fill="none" font-size="10" font-family="sans-serif" text-anchor="middle"><path class="domain" stroke="#000" d="M0.5,6V0.5H1000.5V6"></path><g class="tick" opacity="1" transform="translate(0.5,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="font-family: Arial; text-anchor: end; font-size: 13px;">0 zł</text></g><g class="tick" opacity="1" transform="translate(78.86014164379195,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="font-family: Arial; text-anchor: end; font-size: 13px;">50000 zł</text></g><g class="tick" opacity="1" transform="translate(157.2202832875839,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="font-family: Arial; text-anchor: end; font-size: 13px;">100000 zł</text></g><g class="tick" opacity="1" transform="translate(235.58042493137586,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="font-family: Arial; text-anchor: end; font-size: 13px;">150000 zł</text></g><g class="tick" opacity="1" transform="translate(313.9405665751678,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="font-family: Arial; text-anchor: end; font-size: 13px;">200000 zł</text></g><g class="tick" opacity="1" transform="translate(392.30070821895976,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="font-family: Arial; text-anchor: end; font-size: 13px;">250000 zł</text></g><g class="tick" opacity="1" transform="translate(470.6608498627517,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="font-family: Arial; text-anchor: end; font-size: 13px;">300000 zł</text></g><g class="tick" opacity="1" transform="translate(549.0209915065436,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="font-family: Arial; text-anchor: end; font-size: 13px;">350000 zł</text></g><g class="tick" opacity="1" transform="translate(627.3811331503356,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="font-family: Arial; text-anchor: end; font-size: 13px;">400000 zł</text></g><g class="tick" opacity="1" transform="translate(705.7412747941275,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="font-family: Arial; text-anchor: end; font-size: 13px;">450000 zł</text></g><g class="tick" opacity="1" transform="translate(784.1014164379195,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="font-family: Arial; text-anchor: end; font-size: 13px;">500000 zł</text></g><g class="tick" opacity="1" transform="translate(862.4615580817115,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="font-family: Arial; text-anchor: end; font-size: 13px;">550000 zł</text></g><g class="tick" opacity="1" transform="translate(940.8216997255034,0)"><line stroke="#000" y2="6"></line><text fill="#000" y="9" dy="0.71em" transform="translate(-10,0)rotate(-45)" style="font-family: Arial; text-anchor: end; font-size: 13px;">600000 zł</text></g></g><rect class="bar" x="0" y="186.66666666666666" height="79.99999999999999" width="1000" style="fill: red;"></rect><rect class="bar" x="0" y="53.33333333333334" height="79.99999999999999" width="26.36354874275068" style="fill: red;"></rect><g class="yAxis" fill="none" font-size="10" font-family="sans-serif" text-anchor="end" style="font-size: 13px;"><path class="domain" stroke="#000" d="M-6,320.5H0.5V0.5H-6"></path><g class="tick" opacity="1" transform="translate(0,226.66666666666666)"><line stroke="#000" x2="-6"></line><text fill="#000" x="-9" dy="0.32em" transform="translate(25,0)" font-size="19" rotate="0" style="font-family: Arial; text-anchor: start;">odsetki z tytułu wiboru zmiennego</text></g><g class="tick" opacity="1" transform="translate(0,93.33333333333334)"><line stroke="#000" x2="-6"></line><text fill="#000" x="-9" dy="0.32em" transform="translate(25,0)" font-size="19" rotate="0" style="font-family: Arial; text-anchor: start;">odsetki z tytułu wiboru stałego</text></g></g></g></svg>"""
    svg_text = """
    <svg xmlns="http://www.w3.org/2000/svg" width="100" height="100">
    <text x="50" y="50" font-family="Arial" font-size="20">ąćęłńóśźż</text>
    </svg>
    """
    # convert the SVG string to a byte stream

    
    # convert the SVG byte stream to PNG byte stream
    #drawing = svg2rlg(text_base64)
    drawing = svg2rlg(StringIO(dane_svg))
    png_byte_stream = BytesIO()
    renderPM.drawToFile(drawing, png_byte_stream, fmt="PNG")
    # png_data = renderPM.drawToString(drawing, fmt="PNG")

    # with open('output.png', 'wb') as f:
    #     f.write(png_data)

    # add the image to the document
    document.add_picture(png_byte_stream, width=Inches(7))

    # png_base64 = base64.b64encode(ss)    
    # shape = document.add_picture(png_base64)

    document.add_heading('Dane o kredycie', level=3)

    document.add_paragraph(
        f"Kwota kredytu: {round(float(dane_kredyt['kapital1']),2):,.2f} zł", style='List Bullet'
    )
    document.add_paragraph(
        f"Data uruchomienia kredytu: {dane_kredyt['dataStart1']} " , style='List Bullet'
    )
    document.add_paragraph(
        f"Liczba miesięcy kredytowania: {dane_kredyt['okresy']}" , style='List Bullet'
    )

    document.add_paragraph(
        f"Marża banku: {dane_kredyt['marza']}%" , style='List Bullet'
    )

    document.add_paragraph(
        f"Rodzaj wiboru: Wibor {dane_kredyt['rodzajWiboru']}" , style='List Bullet'
    )

    document.add_paragraph(
        f"Wibor z dnia podpisania umowy: {dane['fin_data']['wibor_start']}%" , style='List Bullet'
    )

    document.add_heading('Zamrożenie wiboru', level=3)

    document.add_paragraph(
        f"Data zamrożenia wiboru: {dane_kredyt['dataZamrozenia']}" , style='List Bullet'
    )

    document.add_paragraph(
        f"Do zwrotu przez bank (raty przed zamrożeniem): {dane['do_zwrotu_przed']}" , style='List Bullet'
    )

    
    document.add_paragraph(
        f"Saldo kredytu: {dane['kapital_do_splaty_po_zamr']}" , style='List Bullet'
    )

    document.add_heading('Zestawienie rat', level=3)
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Dzień'
    hdr_cells[1].text = 'Odsetki'
    hdr_cells[2].text = 'Rata'
    hdr_cells[3].text = "Kapitał do spłaty"
    for d in dane['fin_data']['dane']['raty']:
        row_cells = table.add_row().cells
        row_cells[0].text = d['dzien'][:10]
        row_cells[1].text = d['odsetki'] + " zł"
        row_cells[2].text = d['rata'] + " zł"
        row_cells[3].text = str(d['K_po']) + " zł"
    return document